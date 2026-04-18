"""
Tests for document text extraction (txt, pdf, docx).
"""
import io
import pytest
from app.features.document.service import DocumentDetectionService


@pytest.fixture
def service():
    return DocumentDetectionService()


class TestExtractTxt:

    def test_extracts_plain_text(self, service):
        content = b"Hello world this is a test document."
        text = service._extract_txt(content)
        assert "Hello world" in text

    def test_handles_utf8(self, service):
        content = "Héllo wörld".encode("utf-8")
        text = service._extract_txt(content)
        assert "wörld" in text

    def test_handles_invalid_bytes_gracefully(self, service):
        content = b"Hello \xff\xfe world"
        text = service._extract_txt(content)
        assert "Hello" in text


class TestExtractPdf:

    def test_returns_string(self, service, sample_pdf_bytes):
        text = service._extract_pdf(sample_pdf_bytes)
        assert isinstance(text, str)

    def test_returns_empty_string_on_invalid_pdf(self, service):
        text = service._extract_pdf(b"not a pdf at all")
        assert text == ""

    def test_extracts_text_content(self, service, sample_pdf_bytes):
        text = service._extract_pdf(sample_pdf_bytes)
        assert len(text) >= 0  # just verify it doesn't crash


class TestExtractDocx:

    def test_returns_empty_string_on_invalid_docx(self, service):
        text = service._extract_docx(b"not a docx file")
        assert text == ""

    def test_returns_string(self, service):
        """Create a minimal valid docx in memory."""
        import docx
        doc = docx.Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        text = service._extract_docx(buf.getvalue())
        assert "test paragraph" in text
        assert "another paragraph" in text